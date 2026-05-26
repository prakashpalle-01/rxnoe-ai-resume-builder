from io import BytesIO
from html import escape
import re
from typing import Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer


def build_pdf(resume: dict, template_id: str = "ats-classic") -> bytes:
    config = _template_config(template_id)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=config["margin"],
        leftMargin=config["margin"],
        topMargin=config["top_margin"],
        bottomMargin=config["top_margin"],
    )
    styles = _pdf_styles(config)
    story = []
    p = resume.get("personal_info", {})
    story.append(Paragraph(f"<b>{_safe(p.get('name') or 'Resume')}</b>", styles["ResumeName"]))
    if resume.get("target_title"):
        story.append(Paragraph(f"<b>{_safe(resume.get('target_title'))}</b>", styles["ResumeTitle"]))
    contact = " | ".join(filter(None, [p.get("email"), p.get("phone"), p.get("location"), p.get("linkedin"), p.get("github"), p.get("portfolio")]))
    story.append(Paragraph(_safe(contact), styles["Contact"]))
    story.append(Spacer(1, 7))

    _pdf_section(story, styles, "Summary")
    story.append(Paragraph(_safe(resume.get("summary", "")), styles["Body"]))

    skill_lines = _skill_lines(resume)
    if skill_lines:
        _pdf_section(story, styles, "Technical Skills")
        for label, values in skill_lines:
            story.append(Paragraph(f"<b>{_safe(label)}:</b> {_safe(', '.join(values))}", styles["SkillLine"]))

    if resume.get("experience"):
        bullet_counter = 0
        for index, job in enumerate(resume.get("experience", [])):
            entry = []
            title = job.get("title") or "Role"
            company = job.get("company", "")
            dates = " - ".join(filter(None, [job.get("start_date"), job.get("end_date")]))
            header = f"<b>{_safe(title)}{', ' if company else ''}{_safe(company)}</b>"
            if dates:
                header += f" <font color='#475569'>{_safe(dates)}</font>"
            entry.append(Paragraph(header, styles["ItemHeader"]))
            for bullet in job.get("bullets", []):
                entry.append(Paragraph(_bullet_html(bullet, resume, bullet_counter), styles["ResumeBullet"], bulletText="•"))
                bullet_counter += 1
            entry.append(Spacer(1, 3))
            if index == 0:
                _pdf_section(story, styles, "Experience", entry)
            else:
                story.append(KeepTogether(entry))

    if resume.get("projects"):
        bullet_counter = 0
        for index, project in enumerate(resume.get("projects", [])):
            entry = []
            detail = " | ".join(filter(None, [", ".join(project.get("technologies", [])), project.get("url", "")]))
            header = f"<b>{_safe(project.get('name') or 'Project')}</b>"
            if detail:
                header += f" <font color='#475569'>{_safe(detail)}</font>"
            entry.append(Paragraph(header, styles["ItemHeader"]))
            for bullet in project.get("bullets", []):
                entry.append(Paragraph(_bullet_html(bullet, resume, bullet_counter, project=True), styles["ResumeBullet"], bulletText="•"))
                bullet_counter += 1
            entry.append(Spacer(1, 3))
            if index == 0:
                _pdf_section(story, styles, "Projects", entry)
            else:
                story.append(KeepTogether(entry))

    _pdf_simple_list(story, styles, "Education", resume.get("education", []))
    _pdf_simple_list(story, styles, "Certifications", resume.get("certifications", []))
    doc.build(story)
    return buffer.getvalue()


def build_docx(resume: dict, template_id: str = "ats-classic") -> bytes:
    config = _template_config(template_id)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(config["top_margin"])
        section.bottom_margin = Pt(config["top_margin"])
        section.left_margin = Pt(config["margin"])
        section.right_margin = Pt(config["margin"])
    styles = doc.styles
    styles["Normal"].font.name = config["docx_font"]
    styles["Normal"].font.size = Pt(config["docx_size"])
    styles["Normal"].paragraph_format.space_after = Pt(config["paragraph_spacing"])
    styles["Normal"].paragraph_format.line_spacing = config["line_spacing"]
    styles["Title"].font.name = config["docx_font"]
    styles["Title"].font.size = Pt(config["name_size"])
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 1"].font.name = config["docx_font"]
    styles["Heading 1"].font.size = Pt(config["section_size"])
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = _rgb(config["accent"])
    styles["Heading 1"].paragraph_format.space_before = Pt(9)
    styles["Heading 1"].paragraph_format.space_after = Pt(4)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    p = resume.get("personal_info", {})
    name = doc.add_heading(p.get("name") or "Resume", level=0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if resume.get("target_title"):
        role = doc.add_paragraph()
        role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role.add_run(resume.get("target_title")).bold = True
    contact = doc.add_paragraph(" | ".join(filter(None, [p.get("email"), p.get("phone"), p.get("location"), p.get("linkedin"), p.get("github"), p.get("portfolio")])))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in contact.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(71, 85, 105)
    _docx_section(doc, "Summary", [resume.get("summary", "")])
    skills = [f"{label}: {', '.join(values)}" for label, values in _skill_lines(resume)]
    _docx_section(doc, "Technical Skills", skills)
    if resume.get("experience"):
        doc.add_heading("Experience", level=1)
    bullet_counter = 0
    for job in resume.get("experience", []):
        item = doc.add_paragraph()
        item.paragraph_format.space_before = Pt(4)
        item.add_run(f"{job.get('title', '')}{', ' if job.get('company') else ''}{job.get('company', '')}").bold = True
        dates = " - ".join(filter(None, [job.get("start_date"), job.get("end_date")]))
        if dates:
            item.add_run(f"  {dates}")
        for bullet in job.get("bullets", []):
            _docx_bullet(doc, bullet, resume, bullet_counter)
            bullet_counter += 1
    if resume.get("projects"):
        doc.add_heading("Projects", level=1)
    bullet_counter = 0
    for project in resume.get("projects", []):
        detail = " | ".join(filter(None, [", ".join(project.get("technologies", [])), project.get("url", "")]))
        item = doc.add_paragraph()
        item.paragraph_format.space_before = Pt(4)
        item.add_run(project.get("name", "")).bold = True
        if detail:
            item.add_run(f" | {detail}")
        for bullet in project.get("bullets", []):
            _docx_bullet(doc, bullet, resume, bullet_counter, project=True)
            bullet_counter += 1
    _docx_section(doc, "Education", resume.get("education", []))
    _docx_section(doc, "Certifications", resume.get("certifications", []))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _pdf_styles(config: dict) -> dict:
    base = getSampleStyleSheet()
    base.add(ParagraphStyle(
        name="ResumeName",
        parent=base["Title"],
        fontName=config["font_bold"],
        fontSize=config["name_size"],
        leading=config["name_size"] + 3,
        alignment=TA_CENTER,
        spaceAfter=2,
        textColor=colors.HexColor("#0f172a"),
    ))
    base.add(ParagraphStyle(
        name="ResumeTitle",
        parent=base["Normal"],
        fontName=config["font_bold"],
        fontSize=config["title_size"],
        leading=config["title_size"] + 2.5,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#334155"),
        spaceAfter=2,
    ))
    base.add(ParagraphStyle(
        name="Contact",
        parent=base["Normal"],
        fontName=config["font"],
        fontSize=config["contact_size"],
        leading=config["contact_size"] + 2.2,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#475569"),
    ))
    base.add(ParagraphStyle(
        name="SectionHeading",
        parent=base["Normal"],
        fontName=config["font_bold"],
        fontSize=config["section_size"],
        leading=config["section_size"] + 2,
        textColor=colors.HexColor(config["accent"]),
        spaceBefore=7,
        spaceAfter=2,
    ))
    base.add(ParagraphStyle(name="Body", parent=base["Normal"], fontName=config["font"], fontSize=config["body_size"], leading=config["leading"], spaceAfter=config["paragraph_spacing"]))
    base.add(ParagraphStyle(name="SkillLine", parent=base["Normal"], fontName=config["font"], fontSize=config["skill_size"], leading=config["leading"] - 0.4, spaceAfter=1.8))
    base.add(ParagraphStyle(name="ItemHeader", parent=base["Normal"], fontName=config["font"], fontSize=config["body_size"], leading=config["leading"], spaceBefore=3.5, spaceAfter=2))
    base.add(ParagraphStyle(
        name="ResumeBullet",
        parent=base["Normal"],
        fontName=config["font"],
        fontSize=config["bullet_size"],
        leading=config["leading"],
        leftIndent=13,
        firstLineIndent=-7,
        bulletIndent=2,
        spaceAfter=1.8,
    ))
    return base


def _pdf_section(story, styles, title: str, keep_with: Optional[list] = None) -> None:
    heading = [
        Spacer(1, 4),
        Paragraph(f"<b>{_safe(title.upper())}</b>", styles["SectionHeading"]),
        HRFlowable(width="100%", thickness=0.45, color=colors.HexColor("#cbd5e1"), spaceBefore=0, spaceAfter=4),
    ]
    if keep_with:
        story.append(KeepTogether(heading + keep_with))
    else:
        story.extend(heading)


def _pdf_simple_list(story, styles, title: str, lines: list[str]) -> None:
    lines = [line for line in lines if line]
    if not lines:
        return
    _pdf_section(story, styles, title, [Paragraph(_safe(lines[0]), styles["Body"])])
    for line in lines[1:]:
        story.append(Paragraph(_safe(line), styles["Body"]))


def _skill_lines(resume: dict) -> list[tuple[str, list[str]]]:
    ordered = []
    for group, values in resume.get("skills", {}).items():
        clean_values = [value for value in values if value]
        if clean_values:
            ordered.append((_skill_label(group), clean_values))
    return ordered


def _safe(value: str) -> str:
    return escape(str(value or ""), quote=False)


def _rich_safe(value: str) -> str:
    parts = re.split(r"(\*\*[^*]+\*\*)", str(value or ""))
    rendered = []
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            rendered.append(f"<b>{escape(part[2:-2], quote=False)}</b>")
        else:
            rendered.append(escape(part, quote=False))
    return "".join(rendered)


def _bullet_html(value: str, resume: dict, index: int, project: bool = False) -> str:
    text = str(value or "")
    if "**" in text:
        return _rich_safe(text)
    phrase = _highlight_phrase(text, resume, index, project)
    if not phrase:
        return _safe(text)
    location = text.lower().find(phrase.lower())
    if location < 0:
        return _safe(text)
    return "".join([
        _safe(text[:location]),
        f"<b>{_safe(text[location:location + len(phrase)])}</b>",
        _safe(text[location + len(phrase):]),
    ])


def _highlight_phrase(text: str, resume: dict, index: int, project: bool = False) -> str:
    phrase = _metric_phrase(text)
    if not phrase and (index % 2 == 0 or project):
        phrase = _keyword_phrase(text, resume.get("target_keywords", []))
    return phrase


def _metric_phrase(text: str) -> str:
    match = re.search(r"\b(?:reduced|improved|increased|cut|shortened|accelerated|processed|saved|lowered|raised)[^.;,]*?\b\d+%|\b\d+[%x]\b", text, re.I)
    return match.group(0) if match else ""


def _keyword_phrase(text: str, keywords: list[str]) -> str:
    blocked = {
        "associate", "stack", "product", "skills", "required", "preferred", "development",
        "software", "systems", "team", "business", "technical", "technology", "platform",
        "solution", "engineer", "engineering", "developer", "role", "work", "working",
        "collaboration", "communication", "stakeholder",
    }
    priority = [
        "Ruby on Rails", "React", "Stripe APIs", "Stripe", "PostgreSQL", "Datadog",
        "AWS Performance Insights", "AWS", "GitHub", "customer-data", "admin reporting",
        "donor", "donation", "payment", "Rails",
    ]
    candidates = priority + [keyword for keyword in keywords if keyword]
    seen = set()
    for keyword in sorted(candidates, key=len, reverse=True):
        clean = keyword.strip()
        key = clean.lower()
        if not clean or key in seen or key in blocked or len(clean) < 4:
            continue
        seen.add(key)
        pattern = re.compile(rf"(?<![A-Za-z0-9+#./-]){re.escape(clean)}(?![A-Za-z0-9+#./-])", re.I)
        match = pattern.search(text)
        if match:
            return text[match.start():match.end()]
    return ""


def _docx_section(doc: Document, title: str, lines: list[str]) -> None:
    lines = [line for line in lines if line]
    if not lines:
        return
    doc.add_heading(title, level=1)
    for line in lines:
        doc.add_paragraph(line)


def _docx_bullet(doc: Document, value: str, resume: dict, index: int, project: bool = False) -> None:
    text = str(value or "")
    paragraph = doc.add_paragraph(style="List Bullet")
    phrase = _highlight_phrase(text, resume, index, project)
    if not phrase:
        paragraph.add_run(text)
        return
    location = text.lower().find(phrase.lower())
    paragraph.add_run(text[:location])
    paragraph.add_run(text[location:location + len(phrase)]).bold = True
    paragraph.add_run(text[location + len(phrase):])


def _template_config(template_id: str) -> dict:
    templates = {
        "genai-model": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#111827", "margin": 43, "top_margin": 34, "name_size": 18, "title_size": 10.5, "section_size": 9.5, "body_size": 9.25, "bullet_size": 9, "skill_size": 9, "contact_size": 8.6, "leading": 11.8, "paragraph_spacing": 3, "docx_size": 9.5, "line_spacing": 1.08},
        "ats-classic": {"font": "Times-Roman", "font_bold": "Times-Bold", "docx_font": "Times New Roman", "accent": "#0f172a", "margin": 50, "top_margin": 42, "name_size": 19, "title_size": 11, "section_size": 10.5, "body_size": 10, "bullet_size": 9.6, "skill_size": 9.7, "contact_size": 9, "leading": 12.8, "paragraph_spacing": 4, "docx_size": 10, "line_spacing": 1.12},
        "tech-compact": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#111827", "margin": 40, "top_margin": 30, "name_size": 17, "title_size": 10, "section_size": 9.2, "body_size": 8.9, "bullet_size": 8.7, "skill_size": 8.7, "contact_size": 8.2, "leading": 10.8, "paragraph_spacing": 2, "docx_size": 9, "line_spacing": 1.0},
        "executive-clean": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#1e293b", "margin": 54, "top_margin": 46, "name_size": 20, "title_size": 11, "section_size": 10, "body_size": 9.8, "bullet_size": 9.5, "skill_size": 9.4, "contact_size": 9, "leading": 13.5, "paragraph_spacing": 4, "docx_size": 10, "line_spacing": 1.16},
        "engineer-dense": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#111827", "margin": 39, "top_margin": 29, "name_size": 17, "title_size": 10, "section_size": 9.2, "body_size": 8.8, "bullet_size": 8.6, "skill_size": 8.7, "contact_size": 8.2, "leading": 10.6, "paragraph_spacing": 2, "docx_size": 9, "line_spacing": 1.0},
        "data-modern": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#0891b2", "margin": 47, "top_margin": 38, "name_size": 18, "title_size": 10.5, "section_size": 9.8, "body_size": 9.3, "bullet_size": 9, "skill_size": 9, "contact_size": 8.5, "leading": 12.1, "paragraph_spacing": 3, "docx_size": 9.5, "line_spacing": 1.08},
        "cloud-systems": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#047857", "margin": 47, "top_margin": 38, "name_size": 18, "title_size": 10.5, "section_size": 9.8, "body_size": 9.3, "bullet_size": 9, "skill_size": 9, "contact_size": 8.5, "leading": 12.1, "paragraph_spacing": 3, "docx_size": 9.5, "line_spacing": 1.08},
        "frontend-sharp": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#1d4ed8", "margin": 47, "top_margin": 38, "name_size": 18, "title_size": 10.5, "section_size": 9.8, "body_size": 9.3, "bullet_size": 9, "skill_size": 9, "contact_size": 8.5, "leading": 12.1, "paragraph_spacing": 3, "docx_size": 9.5, "line_spacing": 1.08},
        "minimal-two-page": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#334155", "margin": 56, "top_margin": 48, "name_size": 20, "title_size": 11, "section_size": 10.5, "body_size": 10.2, "bullet_size": 10, "skill_size": 9.9, "contact_size": 9, "leading": 14.2, "paragraph_spacing": 5, "docx_size": 10.5, "line_spacing": 1.18},
        "recruiter-scan": {"font": "Helvetica", "font_bold": "Helvetica-Bold", "docx_font": "Arial", "accent": "#1d4ed8", "margin": 48, "top_margin": 38, "name_size": 19, "title_size": 11, "section_size": 10, "body_size": 9.6, "bullet_size": 9.3, "skill_size": 9.2, "contact_size": 8.8, "leading": 12.6, "paragraph_spacing": 3, "docx_size": 10, "line_spacing": 1.1},
    }
    return templates.get(template_id, templates["ats-classic"])


def _rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _skill_label(value: str) -> str:
    labels = {
        "ai_ml_core": "AI / ML Core",
        "deep_learning": "Deep Learning",
        "genai_llm_systems": "GenAI & Advanced LLM Systems",
        "frameworks_libraries": "Frameworks & Libraries",
        "mlops_engineering": "MLOps & ML Engineering",
        "cloud_infrastructure": "Cloud & Infrastructure",
        "databases_vector_stores": "Databases & Vector Stores",
        "programming": "Programming",
        "monitoring_observability": "Monitoring & Observability",
        "ai_safety_compliance": "AI Safety, Privacy & Compliance",
        "developer_tools": "Developer & Productivity Tools",
        "technical": "Additional Technical Skills",
        "tools": "Tools",
        "cloud": "Cloud",
        "databases": "Databases",
        "soft_skills": "Professional Skills",
    }
    return labels.get(value, value.replace("_", " ").title())
